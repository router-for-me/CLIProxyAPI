from docx import Document

doc = Document('用户协议.docx')
full_text = []

for para in doc.paragraphs:
    if para.text.strip():
        full_text.append(para.text)

# 也读取表格内容
for table in doc.tables:
    for row in table.rows:
        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if row_text:
            full_text.append(' | '.join(row_text))

with open('用户协议_内容.txt', 'w', encoding='utf-8') as f:
    f.write('\n\n'.join(full_text))

print('文档内容已提取到: 用户协议_内容.txt')
print(f'共提取 {len(full_text)} 段内容')
